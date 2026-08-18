from fastapi import UploadFile
import PyPDF2
from docx import Document
from openpyxl import load_workbook
import io
import re
import hashlib
import os
from pathlib import Path
import json

# OCR imports (optional, for PDF scan)
try:
    from pdf2image import convert_from_bytes
    import pytesseract
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False
    print("Warning: OCR libraries not available. PDF scan files won't be readable.")

# Cache directory
if os.environ.get('VERCEL'):
    CACHE_DIR = Path('/tmp/cache')
else:
    CACHE_DIR = Path(__file__).parent / "cache"

try:
    CACHE_DIR.mkdir(exist_ok=True)
except OSError:
    pass

def get_file_hash(content: bytes) -> str:
    """Generate hash for file content to use as cache key"""
    return hashlib.md5(content).hexdigest()

def get_cached_text(file_hash: str) -> str | None:
    """Get cached text if exists"""
    cache_file = CACHE_DIR / f"{file_hash}.txt"
    if cache_file.exists():
        return cache_file.read_text(encoding='utf-8')
    return None

def cache_text(file_hash: str, text: str) -> None:
    """Cache extracted text"""
    cache_file = CACHE_DIR / f"{file_hash}.txt"
    cache_file.write_text(text, encoding='utf-8')

def preprocess_text(text: str) -> str:
    """Clean and normalize extracted text"""
    # Remove excessive whitespace
    text = re.sub(r'\s+', ' ', text)
    # Remove special characters that might confuse AI
    text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f]', '', text)
    # Normalize line breaks
    text = re.sub(r'\n\s*\n', '\n\n', text)
    # Remove empty lines
    lines = [line.strip() for line in text.split('\n') if line.strip()]
    return '\n'.join(lines)

async def extract_text_from_file(file: UploadFile) -> str:
    """Main function to extract text from uploaded file with caching"""
    filename = file.filename.lower()
    content = await file.read()
    
    # Check cache first
    file_hash = get_file_hash(content)
    cached_text = get_cached_text(file_hash)
    if cached_text:
        safe_filename = filename.encode('utf-8', 'ignore').decode('utf-8')
        print(f"Cache hit for {safe_filename}")
        return cached_text
    
    # Extract text based on file type
    if filename.endswith('.pdf'):
        text = extract_from_pdf(content)
    elif filename.endswith('.docx') or filename.endswith('.doc'):
        text = extract_from_docx(content)
    elif filename.endswith('.xlsx') or filename.endswith('.xls'):
        text = extract_from_excel(content)
    elif filename.endswith('.txt'):
        text = content.decode('utf-8')
    else:
        raise ValueError(f"Unsupported file type: {filename}")
    
    # Preprocess text
    text = preprocess_text(text)
    
    # Cache result
    cache_text(file_hash, text)
    
    return text

def extract_from_pdf(content: bytes) -> str:
    """Extract text from PDF, with OCR fallback for scanned PDFs"""
    pdf = PyPDF2.PdfReader(io.BytesIO(content))
    text = ""
    
    # Try normal text extraction first
    for page in pdf.pages:
        page_text = page.extract_text()
        if page_text and page_text.strip():
            text += page_text + "\n"
    
    # If no text found, try OCR (for scanned PDFs)
    if len(text.strip()) < 100 and OCR_AVAILABLE:
        print("No text found in PDF, attempting OCR...")
        try:
            images = convert_from_bytes(content)
            for img in images:
                page_text = pytesseract.image_to_string(img, lang='vie+eng')
                text += page_text + "\n"
        except Exception as e:
            print(f"OCR failed: {e}")
    
    return text.strip()

def extract_from_docx(content: bytes) -> str:
    """Extract text from DOCX including tables"""
    doc = Document(io.BytesIO(content))
    text = []
    
    # Extract paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            text.append(para.text)
    
    # Extract tables
    for table in doc.tables:
        text.append("\n[TABLE]")
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_text:
                text.append(" | ".join(row_text))
        text.append("[/TABLE]\n")
    
    return "\n".join(text)

def extract_from_excel(content: bytes) -> str:
    """Extract text from Excel files"""
    wb = load_workbook(io.BytesIO(content), data_only=True)
    text = []
    
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        text.append(f"\n[SHEET: {sheet_name}]")
        
        for row in ws.iter_rows(values_only=True):
            row_text = [str(cell) if cell is not None else "" for cell in row]
            if any(row_text):
                text.append(" | ".join(row_text))
        
        text.append("[/SHEET]\n")
    
    return "\n".join(text)

def text_to_structured_json(text: str, doc_type: str) -> dict:
    """Convert extracted text to structured JSON for faster AI analysis"""
    # Basic structure extraction
    sections = []
    current_section = {"title": "General", "content": []}
    
    for line in text.split('\n'):
        line = line.strip()
        if not line:
            continue
        
        # Detect section headers (all caps, or starts with number)
        if line.isupper() or re.match(r'^\d+[\.\)]', line):
            if current_section["content"]:
                sections.append(current_section)
            current_section = {"title": line, "content": []}
        else:
            current_section["content"].append(line)
    
    if current_section["content"]:
        sections.append(current_section)
    
    return {
        "document_type": doc_type,
        "total_length": len(text),
        "sections": sections,
        "full_text": text
    }
